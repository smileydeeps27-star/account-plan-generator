#!/usr/bin/env python3
"""
update-trackers.py

Parses generated Word account plans under
  output/FY27 Account Plans/<CP>/<Type>/<Company> - Account Plan.docx
and writes results into:
  - Batch Status Tracker.xlsx     (one row per account)
  - Executive Review Tracker.xlsx (multi-row per account: actions + stakeholders)

Run from the repo root:
    python3 output/update-trackers.py

Override the batch number when needed:
    BATCH_NUM=2 python3 output/update-trackers.py
"""

import os
import sys
from datetime import datetime

from docx import Document
import openpyxl

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT_BASE = os.path.join(REPO_ROOT, "output", "FY27 Account Plans")
BATCH_TRACKER = os.path.join(OUTPUT_BASE, "Batch Status Tracker.xlsx")
EXEC_TRACKER = os.path.join(OUTPUT_BASE, "Executive Review Tracker.xlsx")

BATCH_NUM = int(os.environ.get("BATCH_NUM", "1"))

# Map of "generated docx filename stem" -> "tracker Account column value".
# Add entries when an account's filename differs from the tracker name.
NAME_MAP = {
    "HP Inc.": "HP",
    "Kimberly-Clark": "Kimberly-Clark Corp. (KCC)",
    "Mary Kay Inc.": "MaryKay",
    "McLane Company": "McLane",
    "Hewlett Packard Enterprise": "HPE",
    "Carlsberg Group": "Carlsberg / Britvic",
}

# Names known to be intentionally absent from the tracker (e.g. smoke tests)
SKIP_UNMATCHED = {"Unilever"}


# ---------------------------------------------------------------------------
# Name normalization helper
# ---------------------------------------------------------------------------

def norm(s):
    if s is None:
        return ""
    return str(s).strip().lower().replace("  ", " ")


def tracker_name_for(gen_name, tracker_names):
    """Resolve a generated docx stem to a tracker Account string.

    1. Explicit NAME_MAP hit.
    2. Normalized exact match.
    3. Normalized startswith / substring match (handles trailing " Corp.", etc.).
    """
    if gen_name in NAME_MAP:
        return NAME_MAP[gen_name]
    g = norm(gen_name)
    for t in tracker_names:
        if norm(t) == g:
            return t
    for t in tracker_names:
        tn = norm(t)
        if tn.startswith(g) or g.startswith(tn) or g in tn or tn in g:
            return t
    return None


# ---------------------------------------------------------------------------
# DOCX parser
# ---------------------------------------------------------------------------

def parse_plan_docx(fpath):
    doc = Document(fpath)
    d = {
        "industry": "",
        "revenue": "",
        "employees": "",
        "stakeholders": [],
        "competitors": [],
        "actions": [],
        "exec_pitch": "",
        "why_now": "",
        "entry_point": "",
        "path": os.path.relpath(fpath, OUTPUT_BASE),
    }

    for table in doc.tables:
        if not table.rows:
            continue
        header = [c.text for c in table.rows[0].cells]
        if not header:
            continue

        if header[0] in ("Industry", "Headquarters"):
            for r in table.rows:
                label = r.cells[0].text.strip()
                value = r.cells[1].text.strip() if len(r.cells) > 1 else ""
                if label == "Industry":
                    d["industry"] = value
                elif label == "Annual Revenue":
                    d["revenue"] = value
                elif label == "Employees":
                    d["employees"] = value

        elif "Name" in header and "Role" in header:
            for r in table.rows[1:]:
                cells = [c.text.strip() for c in r.cells]
                if len(cells) >= 2 and cells[0]:
                    d["stakeholders"].append({"name": cells[0], "title": cells[1]})

        elif "Competitor" in header:
            for r in table.rows[1:]:
                name = r.cells[0].text.strip()
                if name:
                    d["competitors"].append(name)

        elif "#" in header and "Action" in header:
            for r in table.rows[1:]:
                cells = [c.text.strip() for c in r.cells]
                if len(cells) >= 3 and cells[1]:
                    d["actions"].append({"action": cells[1], "owner": cells[2]})

    for para in doc.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        if t.startswith("\u201C") and t.endswith("\u201D"):
            d["exec_pitch"] = t[1:-1]
        elif "Why Now:" in t:
            d["why_now"] = t.split("Why Now:", 1)[1].strip()
        elif "Entry Point:" in t:
            d["entry_point"] = t.split("Entry Point:", 1)[1].strip()

    return d


def collect_generated_plans():
    results = {}
    if not os.path.isdir(OUTPUT_BASE):
        print(f"ERROR: {OUTPUT_BASE} not found", file=sys.stderr)
        return results

    for cp_dir in sorted(os.listdir(OUTPUT_BASE)):
        cp_path = os.path.join(OUTPUT_BASE, cp_dir)
        if not os.path.isdir(cp_path) or cp_dir.startswith("."):
            continue
        for type_dir in sorted(os.listdir(cp_path)):
            type_path = os.path.join(cp_path, type_dir)
            if not os.path.isdir(type_path):
                continue
            for f in sorted(os.listdir(type_path)):
                if not f.endswith(".docx"):
                    continue
                company = f.replace(" - Account Plan.docx", "")
                fpath = os.path.join(type_path, f)
                try:
                    results[company] = parse_plan_docx(fpath)
                except Exception as e:
                    print(f"  WARN: failed to parse {fpath}: {e}", file=sys.stderr)
    return results


# ---------------------------------------------------------------------------
# Batch Status Tracker updater
# ---------------------------------------------------------------------------

def update_batch_tracker(generated):
    if not os.path.exists(BATCH_TRACKER):
        print(f"ERROR: {BATCH_TRACKER} not found", file=sys.stderr)
        return 0

    wb = openpyxl.load_workbook(BATCH_TRACKER)
    ws = wb.active

    tracker_names = [
        (row[1].value or "").strip()
        for row in ws.iter_rows(min_row=2)
        if row[1].value
    ]

    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    updated = 0
    unmatched = []

    for gen_name, d in generated.items():
        if gen_name in SKIP_UNMATCHED:
            continue
        mapped = tracker_name_for(gen_name, tracker_names)
        if not mapped:
            unmatched.append(gen_name)
            continue

        for row in ws.iter_rows(min_row=2):
            if (row[1].value or "").strip() == mapped:
                row[8].value = "Complete"
                row[9].value = BATCH_NUM
                row[10].value = now
                row[11].value = d.get("revenue", "")
                row[12].value = d.get("path", "")
                updated += 1
                break

    wb.save(BATCH_TRACKER)

    total_complete = sum(
        1
        for r in openpyxl.load_workbook(BATCH_TRACKER).active.iter_rows(min_row=2, values_only=True)
        if r[8] == "Complete"
    )
    print(f"Batch Status Tracker: {updated} rows updated this run, {total_complete} total Complete")
    if unmatched:
        print(f"  Unmatched ({len(unmatched)}): {', '.join(unmatched)}")
    return updated


# ---------------------------------------------------------------------------
# Executive Review Tracker updater
# ---------------------------------------------------------------------------

def update_exec_tracker(generated):
    if not os.path.exists(EXEC_TRACKER):
        print(f"ERROR: {EXEC_TRACKER} not found", file=sys.stderr)
        return 0

    wb = openpyxl.load_workbook(EXEC_TRACKER)
    ws = wb.active

    tracker_names = sorted({
        (row[1].value or "").strip()
        for row in ws.iter_rows(min_row=2)
        if row[1].value
    })

    accounts_updated = 0
    unmatched = []

    for gen_name, d in generated.items():
        if gen_name in SKIP_UNMATCHED:
            continue
        mapped = tracker_name_for(gen_name, tracker_names)
        if not mapped:
            unmatched.append(gen_name)
            continue

        touched = False
        for row in ws.iter_rows(min_row=2):
            if (row[1].value or "").strip() != mapped:
                continue

            # First row of the account group carries metadata
            if row[0].value is not None:
                row[5].value = d.get("industry", "")
                row[6].value = d.get("revenue", "")
                row[7].value = d.get("employees", "")
                row[8].value = (d.get("exec_pitch") or "")[:500]
                row[9].value = (d.get("why_now") or "")[:300]
                row[10].value = (d.get("entry_point") or "")[:300]
                row[22].value = d["competitors"][0] if d.get("competitors") else ""
                row[24].value = d.get("path", "")
                touched = True

            an = row[11].value
            if isinstance(an, int) and 1 <= an <= len(d.get("actions", [])):
                a = d["actions"][an - 1]
                row[12].value = a["action"]
                row[13].value = a["owner"]
                row[14].value = "Pending"

            sn = row[17].value
            if isinstance(sn, int) and 1 <= sn <= len(d.get("stakeholders", [])):
                s = d["stakeholders"][sn - 1]
                row[18].value = s["name"]
                row[19].value = s["title"]
                row[20].value = "Not Started"

        if touched:
            accounts_updated += 1

    wb.save(EXEC_TRACKER)
    print(f"Executive Review Tracker: {accounts_updated} accounts updated")
    if unmatched:
        print(f"  Unmatched ({len(unmatched)}): {', '.join(unmatched)}")
    return accounts_updated


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    print(f"Scanning plans in: {OUTPUT_BASE}")
    generated = collect_generated_plans()
    print(f"  Parsed {len(generated)} generated plan docx files")

    if not generated:
        print("Nothing to update.")
        return

    update_batch_tracker(generated)
    update_exec_tracker(generated)
    print("Done.")


if __name__ == "__main__":
    main()
